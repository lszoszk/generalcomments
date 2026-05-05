#!/usr/bin/env python3
"""
Jurisprudence ingestion — Phase 1 (CRPD pilot).

Reads OHCHR's bulk jurisprudence dump:
  - catalog.jsonl       (per-case metadata)
  - download_manifest.jsonl (per-file: format, language, file_path, …)
  - raw/en/*.{docx,pdf,doc}

Outputs a paragraph-level corpus + a Tier-1 metadata catalog:
  - json_jurisprudence/<docId>.json    (per case: list of {ID, Section, Labels, Text})
  - jurisprudence_info.json            (one entry per case, no body text)

Schema for each per-case file matches the GC/SP convention plus a `Section`
field holding the court-style section heading ("Facts as submitted by the
author", "The complaint", "Committee's consideration", …).

Schema for `jurisprudence_info.json` is the Tier-1 catalog used by the
website's case browser:

  {
    "docId":        "crpd-c-18-d-22-2014",
    "type":         "jur",
    "treaty":       "CRPD",
    "symbol":       "CRPD/C/18/D/22/2014",
    "country":      "United Republic of Tanzania",
    "year":         2014,
    "title":        "Communication No. 22/2014: Views (Violation found)",
    "outcome":      "violation_found",
    "submittedDate": "23 June 2014",
    "adoptionDate": "18 August 2017",
    "languages":    ["en"],
    "link":         "https://tbinternet.ohchr.org/.../Download.aspx?...",
    "sourceFile":   "json_jurisprudence/crpd-c-18-d-22-2014.json",
    "sourceFormat": "docx",
    "shardId":      "jur_CRPD_2014",   # placeholder — sharder runs separately
    "paragraphCount": 47,
    "wordCount":    8421,
    "labelCount":   38,
    "caseLabels":   ["Persons with disabilities"],
    "firstAddedAt": "2026-04-28",
    "lastVerifiedAt": "2026-04-28"
  }

Only English documents are ingested in v1 (per the user's decision).

Outcome taxonomy (per the user's decision: split violation):
  views                      Substantive merits decision (no specific outcome label)
  violation_found            Committee found a violation
  merits_no_violation        Committee found no violation on the merits
  inadmissible               Communication declared inadmissible
  discontinued               Procedure discontinued (settled, withdrawn, …)
  decision                   Decision (other) — admissibility/procedural ruling
  other                      Anything else (incl. records with no recognisable title)

Usage:
    python3 ingest_jurisprudence.py --treaty CRPD            # pilot — all CRPD
    python3 ingest_jurisprudence.py --treaty CRPD --limit 5  # quick sanity test
    python3 ingest_jurisprudence.py --all                    # full run (~4500 cases)
"""
from __future__ import annotations

import argparse
import html
import json
import re
import shutil
import subprocess
import sys
import tempfile
from datetime import date
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterator
from xml.etree import ElementTree as ET
from zipfile import ZipFile

import fitz  # PyMuPDF — for PDF cases (≈18% of files)

try:
    from docx import Document as DocxDocument
except ImportError:
    print('ERROR: python-docx is required. Install with: pip install python-docx',
          file=sys.stderr)
    sys.exit(1)


W_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
ROOT = Path(__file__).resolve().parent
JURIS_SRC = Path('/Users/lszoszk/Desktop/AI/HURIDOCS/App/output/ohchr_jurisprudence')
CATALOG = JURIS_SRC / 'catalog.jsonl'
MANIFEST = JURIS_SRC / 'download_manifest.jsonl'

OUT_DIR_PARAGRAPHS = ROOT / 'json_jurisprudence'
OUT_INFO = ROOT / 'mysite_pythonanywhere' / 'jurisprudence_info.json'
OCR_DIR = ROOT / 'ocr_jurisprudence'

TREATY_SYMBOL_PREFIXES = {
    'CCPR': ('CCPR/',),
    'CAT': ('CAT/',),
    'CRC': ('CRC/',),
    'CESCR': ('E/C.12/',),
    'CEDAW': ('CEDAW/',),
    'CRPD': ('CRPD/',),
    'CERD': ('CERD/',),
    'CED': ('CED/', 'INT/CED/JUR/'),
}

EXCLUDED_SYMBOLS = {
    # Local OHCHR download resolves to the neighbouring CEDAW/C/84/D/129/2018
    # PDF, so publishing it as 128 would create a false duplicate.
    'CEDAW/C/84/D/128/2018': 'download resolves to CEDAW/C/84/D/129/2018',
}


# ---------------------------------------------------------------------------
# docId slug — `CRPD/C/18/D/22/2014` → `crpd-c-18-d-22-2014`
# ---------------------------------------------------------------------------
def slug(symbol: str) -> str:
    s = symbol.lower().strip()
    s = re.sub(r'[/\.\\:]+', '-', s)
    s = re.sub(r'[^a-z0-9-]', '-', s)
    s = re.sub(r'-+', '-', s).strip('-')
    return s


def symbol_matches_treaty(symbol: str, treaty: str) -> bool:
    """Guard against occasional OHCHR catalog rows filed under the wrong treaty."""
    sym = (symbol or '').strip().upper()
    body = (treaty or '').strip().upper()
    if not sym or not body:
        return True
    expected = TREATY_SYMBOL_PREFIXES.get(body)
    if expected:
        return sym.startswith(expected)
    return True


# ---------------------------------------------------------------------------
# Outcome classifier — title first, body fallback
# ---------------------------------------------------------------------------
def classify_outcome(title: str, body_text: str) -> str:
    t = (title or '').lower()
    b = (body_text or '').lower()
    decision_body = b
    for marker in re.finditer(
        r'\b(?:individual opinion|separate opinion|dissenting opinion)\b',
        b,
    ):
        # Some files mention separate opinions in the header metadata. Only
        # strip them when the marker appears where a trailing opinion section
        # would normally begin.
        if marker.start() > len(b) * 0.6:
            decision_body = b[:marker.start()]
            break
    tail = decision_body[-22000:]

    # Title-based fast path (covers ~70 % of cases)
    if 'inadmissible' in t or 'inadmissibility' in t:
        return 'inadmissible'
    if 'discontinu' in t:
        return 'discontinued'
    if 'no violation' in t or 'merits — no violation' in t:
        return 'merits_no_violation'
    if 'violation' in t:
        # Title says "Views Violation" or "Violation found"
        return 'violation_found'
    if 'views' in t:
        # Generic "Views" — body must clarify
        if ('no violation' in tail or 'has not violated' in tail
            or 'did not violate' in tail
            or 'do not disclose a violation' in tail
            or 'does not disclose a violation' in tail
            or 'do not reveal any violation' in tail
            or 'does not reveal any violation' in tail):
            return 'merits_no_violation'
        if (
            'has failed to fulfil its obligations' in tail
            or 'failed to fulfil its obligations' in tail
            or 'constituted a violation' in tail
            or 'constitute a violation' in tail
            or 'would amount to a breach' in tail
            or 'amounted to a breach' in tail
            or 'failed to discharge its obligations' in tail
            or 'failing to discharge its obligations' in tail
            or re.search(r'(?:considers|finds|concludes)[^.]{0,120}(?:that .{0,60})?violation of (?:article|articles)', tail)
        ):
            return 'violation_found'
        return 'views'
    if 'decision' in t:
        if 'decided to discontinue' in tail or 'decides to discontinue' in tail:
            return 'discontinued'
        if re.search(r'therefore decides:\s*\(?a\)?\s+that the communication is inadmissible', tail):
            return 'inadmissible'
        return 'decision'

    # No useful title — read the decision/conclusion end of the body first.
    if (
        re.search(r'therefore decides:\s*\(?a\)?\s+that the communication is inadmissible', tail)
        or 'declares the communication inadmissible' in tail
        or 'communication is inadmissible under article' in tail
        or 'communication is inadmissible pursuant to article' in tail
    ):
        return 'inadmissible'
    if (
        'decides to discontinue' in tail
        or 'decided to discontinue' in tail
        or 'decision of discontinuance' in tail
        or 'discontinuance decision' in tail
        or 'discontinue the consideration of communication' in tail
    ):
        return 'discontinued'
    if (
        'has not violated' in tail
        or 'did not violate' in tail
        or 'do not disclose a violation' in tail
        or 'does not disclose a violation' in tail
        or 'do not disclose any violation' in tail
        or 'does not disclose any violation' in tail
        or 'do not reveal a violation' in tail
        or 'does not reveal a violation' in tail
        or 'do not reveal any violation' in tail
        or 'does not reveal any violation' in tail
    ):
        return 'merits_no_violation'
    if (
        'has failed to fulfil its obligations' in tail
        or 'failed to fulfil its obligations' in tail
        or 'constituted a violation' in tail
        or 'constitute a violation' in tail
        or 'would amount to a breach' in tail
        or 'amounted to a breach' in tail
        or 'failed to discharge its obligations' in tail
        or 'failing to discharge its obligations' in tail
        or 'has violated the rights' in tail
        or 'has violated her rights' in tail
        or 'has violated his rights' in tail
        or 'has violated their rights' in tail
        or 'infringed the rights' in tail
        or 'would, if implemented, violate' in tail
        or re.search(r'(?:considers|finds|concludes)[^.]{0,120}(?:that .{0,60})?violation of (?:article|articles)', tail)
        or re.search(r'amount(?:s|ed) to a violation of (?:article|articles)', tail)
        or re.search(r'in violation of (?:the author|his|her|their)[^.]{0,120}rights under (?:article|articles)', tail)
        or re.search(r'violated (?:the author|his|her|their|its)[^.]{0,120}rights under article', tail)
        or re.search(r'violated the author[^.]{0,180}rights under (?:article|articles)', tail)
    ):
        return 'violation_found'
    return 'other'


# ---------------------------------------------------------------------------
# Concerned-group labels (v6 patterns from the wider corpus)
# ---------------------------------------------------------------------------
LABEL_PATTERNS = [
    ('Children', [
        r'\bchild(?:ren)?\b', r'\bjuvenile\b', r'\binfant\b', r'\bnewborn\b',
        r'\bminors?\b', r'\bunder.?18\b', r'\bpediatric\b',
        r'\bchild marriage\b', r'\bchild\s+labor(?:ur)?\b',
    ]),
    ('Women/girls', [
        r'\bwom(?:an|en)\b', r'\bgirls?\b', r'\bfemale\b', r'\bmaternal\b',
        r'\bpregnan', r'\bmaternity\b', r'\bmothers?\b', r'\bwidow\b',
        r'\bgender.based violence\b', r'\bgender equality\b',
        r'\bFGM\b', r'\bfemale genital\b', r'\bsexual and reproductive\b',
        r'\bdomestic violence\b', r'\btraffick(?:ing|ed)', r'\bsexual exploit',
    ]),
    ('Persons with disabilities', [
        r'\bdisabilit(?:y|ies)\b', r'\bhandicap\b', r'\bimpairment\b',
        r'\bmental(?:ly)?\s+(?:ill\b|disorder\b|illness\b|health condition)',
        r'\bpsychiatric\b', r'\bcognitive disab\b', r'\bintellectual disab\b',
        r'\breason(?:able)? accommodat', r'\bdeaf(?:ness)?\b',
        r'\bblind(?:ness)?\b', r'\bwheelchair\b', r'\bmental health\b',
        r'\balbinism\b',
    ]),
    ('Migrants', [
        r'\bmigrant\b', r'\bimmigrant\b', r'\blabou?r migration\b',
        r'\bforeign worker\b', r'\bremittance\b',
        r'\bundocumented (?:person|worker|migrant)\b', r'\birregular migra\b',
        r'\bxenophobia\b',
    ]),
    ('Indigenous peoples', [
        r'\bindigenous (?:people|communit|right|land|culture|knowledge|group|person|woman|child)\b',
        r'\btribal (?:people|communit|right|land)\b',
        r'\bfree,?\s*prior\s*and\s*informed\s*consent\b', r'\bFPIC\b',
    ]),
    ('Persons deprived of their liberty', [
        r'\bprison(?:ers?|s)\b', r'\bdetain(?:ee|ment|ed\s+person)\b',
        r'\bincarcerat', r'\bimprison', r'\bjail(?:ed)?\b', r'\bremand(?:ed)?\b',
        r'\bconvict(?:ed|s)\b', r'\bplaces of detention\b',
        r'\bpersons? deprived of (?:their|his|her) liberty\b',
    ]),
    ('Refugees & asylum-seekers', [
        r'\brefugee\b', r'\basylum.seeker\b', r'\basylum seeker\b',
        r'\bnon-refoulement\b', r'\bpersecution\b',
    ]),
    ('Adolescents', [
        r'\badolescent\b', r'\bteen(?:ager)?\b',
        r'\byoung people\b', r'\byoung person\b',
    ]),
    ('Persons living in rural/remote areas', [
        r'\brural\s+(?:area|community|population|household|region|setting|dweller)\b',
        r'\bremote\s+area\b',
    ]),
    ('Persons affected by armed conflict', [
        r'\barmed conflict\b', r'\bwar crime\b', r'\boccupied territory\b',
        r'\bhumanitarian law\b', r'\bIHL\b', r'\bhostilities\b',
        r'\bcombatant\b', r'\bforced displacement\b', r'\bpost.conflict\b',
    ]),
    ('Persons living in poverty', [
        r'\bpovert', r'\bindigent\b', r'\bextreme poverty\b',
        r'\bimpoverish', r'\bdestitut',
        r'\bpoor\s+(?:people|persons|communities|families)\b',
    ]),
    ('Internally displaced persons', [
        r'\bIDPs?\b', r'\binternally displaced\b', r'\binternal displacement\b',
        r'\bforced eviction\b',
    ]),
    ('Persons in street situations', [
        r'\bstreet\s+(?:child|person|people|youth)\b', r'\bhomeless(?:ness)?\b',
    ]),
    ('Children in alternative care', [
        r'\bfoster\s+(?:care|child|parent)\b', r'\borphan',
        r'\bchildren? in (?:alternative|substitute) care\b',
    ]),
    ('Non-citizens and stateless', [
        r'\bstateless(?:ness)?\b', r'\bnon.citizen\b', r'\bnon.nationals?\b',
    ]),
    ('Persons living with HIV/AIDS', [
        r'\bHIV\b', r'\bAIDS\b', r'\bantiretroviral\b',
    ]),
    ('LGBTI+', [
        r'\bLGBT(?:I|Q)?\+?\b', r'\bsexual orientation\b', r'\bgender identity\b',
        r'\bhomosexual(?:ity)?\b', r'\bbisexual\b', r'\btransgender\b',
        r'\bintersex\b', r'\bsame.sex\b',
    ]),
    ('Roma, Gypsies, Sinti and Travellers', [
        r'\bRoma\b', r'\bGyps(?:y|ies)\b', r'\bSinti\b',
    ]),
    ('Persons affected by natural disasters', [
        r'\bnatural disaster\b', r'\bdisaster\s*(?:risk|relief|response|recovery)\b',
        r'\bearthquake\b', r'\bflood(?:ing)?\b', r'\bcyclone\b', r'\bdrought\b',
        r'\bclimate change\b',
    ]),
]
COMPILED_LABELS = [(label, [re.compile(p, re.IGNORECASE) for p in pats])
                   for label, pats in LABEL_PATTERNS]


def label_paragraph(text: str) -> list[str]:
    found = []
    for label, pats in COMPILED_LABELS:
        if any(p.search(text) for p in pats):
            found.append(label)
    return sorted(set(found))


# ---------------------------------------------------------------------------
# DOCX extractor
# ---------------------------------------------------------------------------
PARA_NUM = re.compile(r'^((?:\d+\.\d+(?:\.\d+)*|\d+\.))\s+(.*)', re.DOTALL)
HEADER_KEYWORDS = (
    'committee on the', 'human rights committee', 'committee against torture',
    'views adopted', 'decision adopted', 'communication submitted by:',
    'alleged victim:', 'state party:', 'date of communication:',
    'date of adoption', 'subject matter:', 'procedural issues:',
    'substantive issues:', 'articles of the', 'article of the optional',
)


def _is_non_english_annex_marker(text: str) -> bool:
    """Detect non-English annex blocks that slipped into an English source."""
    t = re.sub(r'\s+', ' ', text.strip())
    return bool(
        re.search(r'\bCOMUNICACI[ÓO]N\s+\(CEDAW\)', t, re.IGNORECASE)
        or re.fullmatch(r'DISPOSICIONES\s+FINALES', t, re.IGNORECASE)
        or re.search(r'\bOriginal\s*:\s*(?:French|Spanish|Arabic|Chinese|Russian)\b', t, re.IGNORECASE)
        or re.search(r'\bOpinion\s+individuelle\b', t, re.IGNORECASE)
    )


def _is_english_annex_marker(text: str) -> bool:
    """Detect a return from a non-English annex to an English annex/opinion."""
    t = re.sub(r'\s+', ' ', text.strip())
    return bool(re.search(r'\bOriginal\s*:\s*English\b', t, re.IGNORECASE))


def _clean_para_id(value: str) -> str:
    return value.rstrip('.') + '.'


def _docx_text_from_el(el: ET.Element) -> str:
    """Extract visible text from a WordprocessingML paragraph/run subtree."""
    parts: list[str] = []
    for node in el.iter():
        if node.tag == f'{W}t' and node.text:
            parts.append(node.text)
        elif node.tag == f'{W}tab':
            parts.append(' ')
        elif node.tag in (f'{W}br', f'{W}cr'):
            parts.append('\n')
    return ''.join(parts)


def _extract_docx_footnote_map(path: Path) -> dict[int, str]:
    """Return DOCX footnotes as {visible_number: text}.

    Word stores true footnotes in ``word/footnotes.xml`` and references them
    from body paragraphs by ``w:footnoteReference/@w:id``. Separator records
    have negative ids and are ignored. For UN jurisprudence files the id is
    the visible footnote number, which matches the website's existing
    ``[[fn:N]]`` marker contract.
    """
    try:
        with ZipFile(path) as z:
            if 'word/footnotes.xml' not in z.namelist():
                return {}
            root = ET.fromstring(z.read('word/footnotes.xml'))
    except Exception:
        return {}

    footnotes: dict[int, str] = {}
    for fn in root.findall('w:footnote', W_NS):
        raw_id = fn.attrib.get(f'{W}id')
        try:
            n = int(raw_id)
        except (TypeError, ValueError):
            continue
        if n < 1:
            continue
        chunks = []
        for p in fn.findall('w:p', W_NS):
            text = re.sub(r'\s+', ' ', _docx_text_from_el(p)).strip()
            if text:
                chunks.append(text)
        text = re.sub(r'\s+', ' ', ' '.join(chunks)).strip()
        if text:
            footnotes[n] = text
    return footnotes


def _iter_docx_body_paragraphs_with_footnote_markers(path: Path) -> list[tuple[str, list[int]]]:
    """Return body paragraph text with ``[[fn:N]]`` tokens in reference order."""
    try:
        with ZipFile(path) as z:
            root = ET.fromstring(z.read('word/document.xml'))
    except Exception:
        doc = DocxDocument(path)
        return [(p.text, []) for p in doc.paragraphs]

    out: list[tuple[str, list[int]]] = []
    body = root.find('w:body', W_NS)
    if body is None:
        return out

    for p in body.findall('w:p', W_NS):
        parts: list[str] = []
        refs: list[int] = []
        for node in p.iter():
            if node.tag == f'{W}t' and node.text:
                parts.append(node.text)
            elif node.tag == f'{W}tab':
                parts.append(' ')
            elif node.tag in (f'{W}br', f'{W}cr'):
                parts.append('\n')
            elif node.tag == f'{W}footnoteReference':
                raw_id = node.attrib.get(f'{W}id')
                try:
                    n = int(raw_id)
                except (TypeError, ValueError):
                    continue
                if n >= 1:
                    parts.append(f'[[fn:{n}]]')
                    refs.append(n)
        out.append((''.join(parts), refs))
    return out


def _footnotes_for_refs(refs: list[int], footnote_map: dict[int, str]) -> list[dict]:
    """Build stable paragraph-local footnote objects from referenced ids."""
    seen: set[int] = set()
    items = []
    for n in refs:
        if n in seen:
            continue
        text = footnote_map.get(n, '').strip()
        if not text:
            continue
        seen.add(n)
        items.append({'n': n, 'text': text})
    return items


def extract_docx_paragraphs(path: Path) -> list[dict]:
    """Walk a jurisprudence DOCX and produce {ID, Section, Labels, Text} records.

    State machine:
      • PRE_BODY  — skip header / metadata block
      • BODY      — collect numbered paragraphs, track section heading via
                    inter-paragraph plain text
    """
    doc = DocxDocument(path)
    footnote_map = _extract_docx_footnote_map(path)
    body_paragraphs = _iter_docx_body_paragraphs_with_footnote_markers(path)
    state = 'PRE_BODY'
    current_section = None
    paragraphs: list[dict] = []
    pending_text: list[str] = []
    pending_refs: list[int] = []
    pending_id: str | None = None
    pre_body_intro: list[str] = []

    def _flush():
        nonlocal pending_refs
        if pending_id and pending_text:
            text = ' '.join(t.strip() for t in pending_text if t.strip())
            text = re.sub(r'\s+', ' ', text).strip()
            if len(text) >= 20:
                row = {
                    'ID': pending_id,
                    'Section': current_section or '',
                    'Labels': [],
                    'Text': text,
                }
                footnotes = _footnotes_for_refs(pending_refs, footnote_map)
                if footnotes:
                    row['Footnotes'] = footnotes
                paragraphs.append(row)
        pending_refs = []

    para_source = body_paragraphs if body_paragraphs else [(p.text, []) for p in doc.paragraphs]
    for raw_text, refs in para_source:
        t = raw_text.strip()
        if not t:
            continue
        if _is_non_english_annex_marker(t):
            if state == 'PRE_BODY':
                continue
            _flush()
            break

        # Identify the leading numbered marker, if any
        m = PARA_NUM.match(t)
        if m and _is_non_english_annex_marker(m.group(2)):
            _flush()
            break

        # Switch to BODY mode the first time we see a "1." or "1.1" leading
        if state == 'PRE_BODY':
            if m and m.group(1).split('.')[0] == '1':
                state = 'BODY'
                if pre_body_intro and not m.group(1).startswith('1'):
                    text = re.sub(r'\s+', ' ', ' '.join(pre_body_intro)).strip()
                    if len(text) >= 80:
                        paragraphs.append({
                            'ID': '1.',
                            'Section': '',
                            'Labels': [],
                            'Text': text,
                        })
                pending_id = _clean_para_id(m.group(1))
                pending_text = [m.group(2)]
                pending_refs = list(refs)
            elif _is_docx_intro_paragraph(t):
                pre_body_intro.append(t)
            elif _is_docx_section_heading(t):
                current_section = t.rstrip('.')
            # otherwise skip (header, metadata)
            continue

        # BODY mode
        if m:
            # New numbered paragraph — flush previous, start new
            _flush()
            pending_id = _clean_para_id(m.group(1))
            pending_text = [m.group(2)]
            pending_refs = list(refs)
        else:
            # Non-numbered line in body. Two cases:
            #  (a) Section heading — short line, capitalised, NOT a continuation
            #  (b) Continuation of the current numbered paragraph
            looks_like_heading = (
                len(t) <= 120
                and not t.endswith(',')
                and not t.endswith(';')
                and not _looks_like_continuation(t)
            )
            if looks_like_heading and not _is_signature_line(t):
                # Treat as section heading. But first flush the open paragraph.
                _flush()
                pending_id = None
                pending_text = []
                pending_refs = []
                current_section = t.rstrip('.')
            else:
                pending_text.append(t)
                pending_refs.extend(refs)
    _flush()
    if len(paragraphs) >= 5:
        return apply_paragraph_namespaces(paragraphs)
    mixed_fallback = extract_docx_mixed_unnumbered_decision(doc)
    if len(mixed_fallback) > len(paragraphs):
        return apply_paragraph_namespaces(mixed_fallback)
    if paragraphs:
        return apply_paragraph_namespaces(paragraphs)
    return apply_paragraph_namespaces(extract_docx_unnumbered_decision(doc))


def extract_docx_mixed_unnumbered_decision(doc: DocxDocument) -> list[dict]:
    """Fallback for legacy DOCX where Word lost early paragraph numbering.

    Some older HRC files keep the body text but drop the automatic list numbers
    until a later section. We preserve official explicit markers when present
    and generate cautious section-local IDs for long unnumbered body paragraphs.
    """
    paragraphs = []
    current_section = ''
    started = False
    intro_count = 0
    facts_count = 0
    generated_count = 0

    def add(pid: str, section: str, text: str, generated: bool = False) -> None:
        row = {
            'ID': pid,
            'Section': section,
            'Labels': [],
            'Text': re.sub(r'\s+', ' ', text).strip(),
        }
        if generated:
            row['GeneratedParagraphID'] = True
            row['GeneratedIDReason'] = 'legacy_docx_missing_numbering'
        paragraphs.append(row)

    for p in doc.paragraphs:
        t = re.sub(r'\s+', ' ', p.text.strip())
        if not t:
            continue
        low = t.lower().strip(' .')
        if _is_non_english_annex_marker(t):
            if started:
                break
            continue
        if not started:
            if (
                low.startswith(('views under article', 'decision under article'))
                or low.startswith(('the first author is', 'the author of the communication is', 'the authors of the communication are'))
            ):
                started = True
                if low.startswith(('views under article', 'decision under article')):
                    current_section = t.rstrip('.')
                    continue
            else:
                continue

        m = PARA_NUM.match(t)
        if m:
            add(_clean_para_id(m.group(1)), current_section, m.group(2))
            continue

        if _is_docx_section_heading(t) or re.match(r'^(?:the case of|the facts as submitted by the authors)\b', t, re.IGNORECASE):
            current_section = t.rstrip('.')
            continue

        if len(t) < 80:
            continue
        if ':' in t[:45] and not low.startswith(('the committee', 'the author', 'the first author', 'the second author')):
            continue
        if low.startswith((
            'the human rights committee, established',
            'having concluded',
            'having taken into account',
            'meeting on ',
            'adopts the following',
        )):
            continue

        section_low = current_section.lower()
        if not paragraphs or section_low.startswith(('views under article', 'decision under article')):
            intro_count += 1
            add(f'1.{intro_count}.', current_section, t, generated=True)
        elif 'facts' in section_low or section_low.startswith(('factual', 'background', 'the case of')):
            facts_count += 1
            add(f'2.{facts_count}.', current_section, t, generated=True)
        else:
            generated_count += 1
            add(f'G{generated_count}.', current_section, t, generated=True)
    return paragraphs


def extract_docx_unnumbered_decision(doc: DocxDocument) -> list[dict]:
    """Fallback for very short discontinuance decisions with no numbered body."""
    paragraphs = []
    source_paragraphs = [re.sub(r'\s+', ' ', p.text.strip()) for p in doc.paragraphs]
    non_empty = [p for p in source_paragraphs if p]
    legacy_short = len(non_empty) <= 12
    for text in source_paragraphs:
        if len(text) < 80:
            if not (legacy_short and text.lower().startswith('the communication is inadmissible')):
                continue
        low = text.lower()
        # Header metadata can be long too ("Substantive issues: ...").
        if ':' in text[:45] and not low.startswith(('at its meeting', 'the committee')):
            continue
        if legacy_short and (
            low.startswith('the author of the communication')
            or low.startswith('article 2 ')
            or low.startswith('being unable to find')
            or low.startswith('the communication is inadmissible')
        ):
            pass
        elif not (
            low.startswith('at its meeting')
            or low.startswith('the committee')
            or low.startswith('decides to discontinue')
            or 'decided to discontinue' in low
            or 'decides to discontinue' in low
            or 'declares the communication inadmissible' in low
        ):
            continue
        paragraphs.append({
            'ID': f'{len(paragraphs) + 1}.',
            'Section': 'Decision',
            'Labels': [],
            'Text': text,
            'GeneratedParagraphID': legacy_short,
            'GeneratedIDReason': 'legacy_unnumbered_decision' if legacy_short else '',
        })
    return paragraphs


# ---------------------------------------------------------------------------
# HTML extractors
# ---------------------------------------------------------------------------

class _JurisHtmlParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.items: list[tuple[str, dict, str]] = []
        self._tag: str | None = None
        self._attrs: dict = {}
        self._buf: list[str] = []

    def handle_starttag(self, tag, attrs):
        if tag.lower() in ('p', 'h1', 'h2', 'h3', 'h4'):
            self._tag = tag.lower()
            self._attrs = dict(attrs)
            self._buf = []

    def handle_data(self, data):
        if self._tag:
            self._buf.append(data)

    def handle_endtag(self, tag):
        if self._tag and tag.lower() == self._tag:
            text = html.unescape(''.join(self._buf))
            text = _clean_html_text(text)
            if text:
                self.items.append((self._tag, self._attrs, text))
            self._tag = None
            self._attrs = {}
            self._buf = []


def _clean_html_text(text: str) -> str:
    replacements = {
        '\x13': '-',
        '\x14': '-',
        '\x18': "'",
        '\x19': "'",
        '\x1c': '"',
        '\x1d': '"',
        '\uf02a': '',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace('\xa0', ' ')
    return re.sub(r'\s+', ' ', text).strip()


def _read_html_text(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ('utf-16', 'utf-8-sig', 'utf-8', 'latin-1'):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode('utf-8', errors='replace')


def _is_html_heading(tag: str, attrs: dict, text: str) -> bool:
    style = (attrs.get('style') or '').lower().replace(' ', '')
    return tag.startswith('h') or 'font-weight:bold' in style or _is_docx_section_heading(text)


def extract_html_paragraphs(path: Path) -> list[dict]:
    """Extract jurisprudence paragraphs from OHCHR DocStore HTML exports."""
    parser = _JurisHtmlParser()
    parser.feed(_read_html_text(path))
    paragraphs: list[dict] = []
    current_section = ''
    started = False
    intro_count = 0
    facts_count = 0
    generated_count = 0

    def add(pid: str, section: str, text: str, generated_reason: str = '') -> None:
        row = {
            'ID': pid,
            'Section': section,
            'Labels': [],
            'Text': _clean_extracted_text(text),
        }
        if generated_reason:
            row['GeneratedParagraphID'] = True
            row['GeneratedIDReason'] = generated_reason
        if len(row['Text']) >= 20:
            paragraphs.append(row)

    for tag, attrs, text in parser.items:
        low = text.lower().strip(' .')
        if not started:
            if low.startswith(('views under article', 'decision under article', 'decision on admissibility')):
                started = True
                current_section = text.rstrip('.')
            continue
        if low.startswith(('[done in', '[adopted in', 'home || treaties')):
            break
        if _is_signature_line(text):
            continue

        m = PARA_NUM.match(text) or re.match(r'^(\d+(?:\.\d+)*\.)\s*(.+)', text)
        if m:
            add(_clean_para_id(m.group(1)), current_section, m.group(2))
            continue

        if _is_html_heading(tag, attrs, text):
            if not low.startswith(('adopts the following', 'views under article', 'decision under article')):
                current_section = text.rstrip('.')
            continue

        if len(text) < 40:
            continue
        if low.startswith((
            'the human rights committee, established',
            'meeting on ',
            'having concluded',
            'having taken into account',
            'adopts the following',
        )):
            continue
        if ':' in text[:45] and not low.startswith(('the author', 'the first author', 'the second author', 'the committee')):
            continue

        section_low = current_section.lower()
        if not paragraphs or section_low.startswith(('views under article', 'decision under article')):
            intro_count += 1
            add(f'1.{intro_count}.', current_section, text, 'html_missing_numbering')
        elif 'facts' in section_low or section_low.startswith(('factual', 'background', 'the case of')):
            facts_count += 1
            add(f'2.{facts_count}.', current_section, text, 'html_missing_numbering')
        else:
            generated_count += 1
            add(f'G{generated_count}.', current_section, text, 'html_missing_numbering')
    return apply_paragraph_namespaces(_repair_html_generated_ids(paragraphs))


def _repair_html_generated_ids(paragraphs: list[dict]) -> list[dict]:
    """Turn generated section IDs into official-looking sequence IDs when clear.

    OHCHR HTML exports sometimes drop the visible numbers for a run of
    paragraphs but resume with e.g. `6.6` or `7.3`. If the generated paragraphs
    immediately precede that explicit ID in the same section, recover `6.1`...
    """
    by_section: dict[str, list[int]] = {}
    for idx, para in enumerate(paragraphs):
        by_section.setdefault(para.get('Section') or '', []).append(idx)
    for indexes in by_section.values():
        pending: list[int] = []
        for idx in indexes:
            para = paragraphs[idx]
            pid = para.get('ID') or ''
            if para.get('GeneratedIDReason') == 'html_missing_numbering' and pid.startswith('G'):
                pending.append(idx)
                continue
            explicit = para_id_tuple(pid)
            if pending and explicit and len(explicit) == 2 and explicit[1] == len(pending) + 1:
                parent = explicit[0]
                for offset, pending_idx in enumerate(pending, start=1):
                    generated = paragraphs[pending_idx]
                    generated['OriginalID'] = generated.get('ID')
                    generated['ID'] = f'{parent}.{offset}.'
                    generated['IdCorrection'] = 'html_sequence_recovered'
                pending = []
            elif explicit:
                pending = []
    return paragraphs


def _looks_like_continuation(text: str) -> bool:
    """Heuristic: if a non-numbered line looks like a sentence fragment, it's
    a continuation of the previous numbered paragraph rather than a heading."""
    if not text:
        return False
    if text[0].islower():
        return True
    # Common continuation starters
    starters = ('and ', 'or ', 'but ', 'however ', 'therefore ', 'in this ',
                'it ', 'they ', 'this ', 'these ', 'such ', 'thus ')
    if text.lower().startswith(starters):
        return True
    return False


def _is_docx_intro_paragraph(text: str) -> bool:
    low = re.sub(r'\s+', ' ', text.strip().lower())
    return (
        len(text) >= 80
        and low.startswith((
            'the author of the communication is',
            'the authors of the communication are',
            'the author is',
            'the authors are',
        ))
    )


def _is_docx_section_heading(text: str) -> bool:
    t = re.sub(r'\s+', ' ', text.strip())
    if not t or len(t) > 130 or ':' in t:
        return False
    if _looks_like_continuation(t) or _is_signature_line(t):
        return False
    low = t.lower().strip(' .')
    return bool(PDF_HEADING_RE.match(low))


def _is_signature_line(text: str) -> bool:
    """Strip off the trailing parenthetical signatures, footnotes, etc."""
    # GE.NN-NNNNN reference numbers
    if re.match(r'^GE\.\d{2}-\d{4,6}', text):
        return True
    if re.match(r'^[A-Z]{2,5}/[A-Z]+/\d+/\w', text):  # doc symbol leak
        return True
    if text.startswith('*'):  # footnote
        return True
    return False


# ---------------------------------------------------------------------------
# PDF + legacy DOC extractors
# ---------------------------------------------------------------------------
PDF_PARA_MARKER = re.compile(r'^(\d{1,3}(?:\.\d+)+\.?|\d{1,3}[.,])(?:\s+(.*)|\s*)$')
PDF_TOPLEVEL_MARKER = re.compile(r'^(\d{1,2})\s+([A-Z][^\t].*)$')
PDF_SECTION_LETTER = re.compile(r'^[A-Z]\.$')
PDF_HEADING_RE = re.compile(
    r'^(?:'
    r'background|'
    r'(?:the\s+)?facts(?:\s+as\s+(?:submitted|presented)\s+by\s+the\s+authors?)?|'
    r'factual\s+background|'
    r'complaints?|'
    r'(?:the\s+)?state\s+party[’\']?s\s+(?:observations|submissions?)|'
    r'(?:the\s+)?author[’\']?s\s+(?:comments|observations)|'
    r'issues?\s+and\s+proceedings\s+before\s+the\s+committee|'
    r'consideration\s+of\s+(?:admissibility|the\s+merits)|'
    r'(?:the\s+)?committee[’\']?s\s+consideration|'
    r'admissibility|merits|conclusions?|'
    r'decision\s+on\s+admissibility|'
    r'individual\s+opinion(?:\s+.*)?|separate\s+opinion(?:\s+.*)?|dissenting\s+opinion(?:\s+.*)?|concurring\s+opinion(?:\s+.*)?|annex'
    r')$',
    re.IGNORECASE,
)
TESSERACT_TSV_LEAK = re.compile(
    r'(?<!\d)[1-5][\t ]+\d+[\t ]+\d+[\t ]+\d+[\t ]+\d+[\t ]+\d+[\t ]+\d+[\t ]+\d+[\t ]+\d+[\t ]+\d+[\t ]+-?\d+(?:\.\d+)?[\t ]+'
)
TESSERACT_TSV_LINE_START = re.compile(
    r'(?m)^[1-5][\t ]+\d+[\t ]+\d+[\t ]+\d+[\t ]+\d+[\t ]+\d+[\t ]+\d+[\t ]+\d+[\t ]+\d+[\t ]+\d+[\t ]+-?\d+(?:\.\d+)?[\t ]+'
)


def normalize_para_id(raw: str) -> str:
    """Return the house paragraph-ID shape: `1.` / `2.1.` / `7.3.2.`."""
    s = raw.strip()
    s = s[:-1] if s.endswith(('.', ',')) else s
    return f'{s}.'


def para_id_tuple(raw: str | None) -> tuple[int, ...] | None:
    if not raw:
        return None
    s = raw.strip()
    s = s[:-1] if s.endswith(('.', ',')) else s
    if not re.fullmatch(r'\d+(?:\.\d+)*', s):
        return None
    return tuple(int(part) for part in s.split('.'))


def _strip_tesseract_tsv_leaks(text: str) -> str:
    """Remove tab-separated Tesseract TSV rows that occasionally leak into OCR.

    The leak usually appears attached to a valid word, e.g.
    ``the5<TSV metadata>only5<TSV metadata>thing``. Replacing the TSV metadata
    with a space preserves the surrounding words: ``the only thing``.
    """
    text = TESSERACT_TSV_LEAK.sub(' ', text)
    text = TESSERACT_TSV_LINE_START.sub('', text)
    return text


def _pdf_marker_action(raw_id: str, rest: str, current_id: str | None) -> str:
    """Return `new`, `append_line`, or `append_rest` for a candidate marker."""
    if not current_id:
        return 'new'
    rest = (rest or '').strip()
    current = para_id_tuple(current_id)
    candidate = para_id_tuple(raw_id)
    if not current or not candidate:
        return 'new'
    if rest.startswith((',', ';', ':', ')', ']')):
        return 'append_line'
    if rest and rest[0].islower():
        return 'append_line'
    if len(candidate) == 1 and candidate[0] >= 40:
        return 'append_rest'
    # OCR sometimes turns continuing body text into a paragraph-looking marker,
    # e.g. "25.4 square miles..." in the middle of paragraph 19.1.
    if rest and rest[0].islower() and candidate[0] > current[0] + 1:
        return 'append_rest'
    # Cited case numbers such as "10.145, 10.305..." or TSV residue can look
    # like impossible paragraph IDs. Keep them in the text, do not split.
    if len(candidate) > 1 and candidate[-1] >= 40:
        return 'append_line'
    return 'new'


def _is_front_matter_date_marker(raw_id: str, rest: str, has_paragraphs: bool, current_section: str) -> bool:
    """Skip OCR front-matter dates that look like top-level paragraphs."""
    if has_paragraphs:
        return False
    if current_section and current_section.strip().lower() != 'annex':
        return False
    candidate = para_id_tuple(raw_id)
    if not candidate or len(candidate) != 1 or candidate[0] <= 1:
        return False
    low = (rest or '').strip().lower()
    return bool(
        (low and low[0].islower())
        or re.match(
            r'^(?:january|february|march|april|may|june|july|august|september|october|november|december)\b',
            low,
        )
        or 'original:' in low
        or 'human rights committee' in low
        or 'submitted by:' in low
    )


def _set_corrected_para_id(para: dict, corrected_id: str, reason: str) -> None:
    if para.get('ID') == corrected_id:
        return
    para.setdefault('RawID', para.get('ID'))
    para['ID'] = corrected_id
    para['IdCorrection'] = reason


def repair_paragraph_id_sequence(paragraphs: list[dict]) -> list[dict]:
    """Correct narrow, sequence-obvious OCR errors in paragraph IDs.

    This does not invent missing paragraphs. It only repairs IDs when adjacent
    numbering makes the intended value clear, e.g. `4.4, 4.2, 4.3` → `4.1,
    4.2, 4.3` or `2.20, 2.11, 2.12` → `2.10, 2.11, 2.12`.
    """
    for _ in range(3):
        changed = False
        for i in range(1, len(paragraphs) - 1):
            prev = paragraphs[i - 1]
            cur = paragraphs[i]
            nxt = paragraphs[i + 1]
            if (prev.get('Section') or '') != (cur.get('Section') or ''):
                continue
            if (cur.get('Section') or '') != (nxt.get('Section') or ''):
                continue
            a = para_id_tuple(prev.get('ID'))
            b = para_id_tuple(cur.get('ID'))
            c = para_id_tuple(nxt.get('ID'))
            if not a or not b or not c:
                continue
            if len(a) == len(b) == len(c) == 2 and a[0] == b[0] == c[0]:
                parent = a[0]
                # Previous ID misread high: 4.4, 4.2, 4.3 -> 4.1, 4.2, 4.3.
                if a[1] > b[1] and c[1] == b[1] + 1 and b[1] > 1:
                    corrected = f'{parent}.{b[1] - 1}.'
                    _set_corrected_para_id(prev, corrected, 'sequence_previous_high')
                    changed = True
                    continue
                # Current ID misread low: 3.7, 3.6, 3.9 -> 3.7, 3.8, 3.9.
                if b[1] < a[1] and c[1] == a[1] + 2:
                    corrected = f'{parent}.{a[1] + 1}.'
                    _set_corrected_para_id(cur, corrected, 'sequence_current_low')
                    changed = True
            # Current parent misread: 5.1, 5.2, 8.3, 5.4 -> 5.1, 5.2, 5.3, 5.4.
            if len(a) == len(b) == len(c) == 2 and a[0] == c[0] and c[1] == a[1] + 2 and b[1] == a[1] + 1:
                corrected = f'{a[0]}.{b[1]}.'
                _set_corrected_para_id(cur, corrected, 'sequence_current_parent')
                changed = True
        if not changed:
            break
    return paragraphs


def apply_paragraph_namespaces(paragraphs: list[dict]) -> list[dict]:
    """Prefix paragraph IDs in official trailing materials with stable namespaces.

    Main-body IDs remain untouched. Separate/individual/dissenting opinions use
    OP1-, OP2-, ...; real appendix/annex materials use A1-, A2-, ... . A bare
    historical "Annex" heading from UN report compilations is not treated as a
    case annex because many old HRC decisions live under that generic heading.
    """
    opinion_count = 0
    annex_count = 0
    current_prefix = ''
    last_section = object()
    opinion_zone = False

    for para in paragraphs:
        section = (para.get('Section') or '').strip()
        low = section.lower()
        section_changed = section != last_section
        if section_changed:
            is_opinion = bool(re.search(r'\b(?:individual|separate|dissenting|concurring)\s+opinion\b', low))
            is_annex = bool(re.search(r'\b(?:appendix|annex\s+(?:\d+|[ivx]+|[a-z]))\b', low, re.IGNORECASE))
            if is_opinion or (opinion_zone and re.match(r'^[A-Z]\.\s+', section)):
                opinion_count += 1
                current_prefix = f'OP{opinion_count}'
                opinion_zone = True
            elif is_annex:
                annex_count += 1
                current_prefix = f'A{annex_count}'
            elif not opinion_zone:
                current_prefix = ''
            last_section = section

        if current_prefix:
            original = para.get('OriginalID') or para.get('ID')
            if original:
                para['OriginalID'] = original
                para['Namespace'] = current_prefix
                para['ID'] = f'{current_prefix}-{original}'
    return paragraphs


def _is_opinion_section(section: str) -> bool:
    low = (section or '').strip().lower()
    return bool(re.search(r'\b(?:individual|separate|dissenting|concurring)\s+opinion\b', low))


def _is_namespaced_continuation_section(section: str, opinion_zone: bool) -> bool:
    return bool(opinion_zone and re.match(r'^[A-Z]\.\s+', (section or '').strip()))


def _is_pdf_noise_line(line: str) -> bool:
    t = line.strip()
    if not t:
        return True
    if re.match(r'^[A-Z]{2,6}/C/\d+/D/\d+/\d+(?:/Rev\.\d+)?$', t):
        return True
    if re.match(r'^page\s+\d+$', t, re.IGNORECASE):
        return True
    if re.match(r'^GE\.\d{2}-\d{4,6}', t):
        return True
    if re.match(r'^[A-Z](?:\.\s*)?[A-Za-z .’\'-]+\s*[\[{]\s*signed\s*[\]}]', t, re.IGNORECASE):
        return True
    if t.startswith('[Done in') or re.match(r'^Subsequently to be (?:issued|translated)', t, re.IGNORECASE):
        return True
    if re.match(r'^Committee[’\']?s Annual Report to the General Assembly\.?\]?$', t, re.IGNORECASE):
        return True
    return False


def _is_pdf_heading(line: str, previous_text: str = '') -> bool:
    """Best-effort heading detector for PDF text streams.

    Jurisprudence PDFs usually put headings as short standalone lines between
    numbered paragraphs. We only split an open paragraph on these when the
    buffered paragraph already ends cleanly, which prevents normal wrapped
    sentence lines from being misread as sections.
    """
    t = line.strip()
    if not t or len(t) > 130:
        return False
    if _is_signature_line(t):
        return False
    if PDF_SECTION_LETTER.match(t):
        return True
    if re.match(r'^(?:(?:A|B|C|D|E|F)\.\s+)?individual\s+opinion\b', t, re.IGNORECASE):
        return True
    if re.match(r'^(?:A|B|C|D|E|F)\.\s+.*\bindividual\s+opinion\b', t, re.IGNORECASE):
        return True
    if t.endswith((',', ';', ':')):
        return False
    if t[0].islower():
        return False
    if previous_text and previous_text.rstrip()[-1:] not in '.!?)”’]':
        return False
    low = re.sub(r'\s+', ' ', t.lower()).strip(' .')
    if PDF_HEADING_RE.match(low):
        return True
    lettered = re.match(r'^(?:A|B|C|D|E|F)\.\s+(.+)', t)
    if lettered:
        tail = lettered.group(1).lower()
        return bool(re.search(
            r'\b(?:decision|opinion|admissibility|merits|facts?|complaints?|communication|proceedings|observations?|comments?|annex)\b',
            tail,
        ))
    return False


def _section_needs_continuation(section: str) -> bool:
    t = (section or '').strip()
    low = t.lower()
    return bool(
        _is_opinion_section(t)
        and (
            t.endswith(',')
            or re.search(r'\bcommittee\s+member$', low)
            or re.search(r'\bcommittee\s+members?$', low)
            or re.search(r'\band\s+[A-Z][A-Za-z’\'-]+$', t)
        )
    )


def _clean_extracted_text(text: str) -> str:
    text = _strip_tesseract_tsv_leaks(text)
    text = re.sub(r'-\s*\n\s*', '', text)
    text = re.sub(r'(?<=[.!?”\)])\s*\d{1,3}(?=\s+[A-Z])', '', text)
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()


def _parse_pdf_text_pages(pages: list[str]) -> list[dict]:
    paragraphs: list[dict] = []
    current_id: str | None = None
    current_section = ''
    pending_heading = ''
    buf: list[str] = []
    unnumbered_buf: list[str] = []
    unnumbered_counter = 0
    opinion_zone = False
    skip_non_english = False

    def flush() -> None:
        nonlocal current_id, buf
        if not current_id or not buf:
            current_id = None
            buf = []
            return
        text = _clean_extracted_text(' '.join(x.strip() for x in buf if x.strip()))
        if len(text) >= 20:
            paragraphs.append({
                'ID': current_id,
                'Section': current_section,
                'Labels': [],
                'Text': text,
            })
        current_id = None
        buf = []

    def flush_unnumbered() -> None:
        nonlocal unnumbered_buf, unnumbered_counter
        if not unnumbered_buf:
            return
        text = _clean_extracted_text(' '.join(x.strip() for x in unnumbered_buf if x.strip()))
        unnumbered_buf = []
        if len(text) < 40:
            return
        unnumbered_counter += 1
        paragraphs.append({
            'ID': f'U{unnumbered_counter}.',
            'Section': current_section,
            'Labels': [],
            'Text': text,
            'GeneratedID': True,
            'GeneratedIDReason': 'unnumbered_opinion_paragraph',
        })

    def starts_new_unnumbered_paragraph(line: str) -> bool:
        if not unnumbered_buf:
            return False
        previous = ' '.join(unnumbered_buf).rstrip()
        if line.startswith(('-', '•')):
            return True
        if previous.endswith(('.', '!', '?', ';', ':', ']', ')')) and line[:1].isupper():
            return True
        return False

    for page_text in pages:
        page_text = _strip_tesseract_tsv_leaks(page_text)
        for raw_line in page_text.splitlines():
            line = raw_line.strip()
            if not line:
                flush_unnumbered()
                continue
            if skip_non_english:
                if _is_english_annex_marker(line):
                    skip_non_english = False
                    current_section = ''
                    opinion_zone = False
                    unnumbered_counter = 0
                continue
            if _is_pdf_noise_line(line):
                continue
            if _is_non_english_annex_marker(line):
                if not paragraphs and not current_id and not current_section:
                    continue
                flush_unnumbered()
                flush()
                skip_non_english = True
                current_section = ''
                opinion_zone = False
                continue

            marker = PDF_PARA_MARKER.match(line)
            if not marker and not opinion_zone:
                top_marker = PDF_TOPLEVEL_MARKER.match(line)
                if top_marker:
                    current_tuple = para_id_tuple(current_id)
                    candidate_top = int(top_marker.group(1))
                    if not current_tuple or candidate_top == current_tuple[0] + 1:
                        marker = top_marker
            if marker:
                raw_id = marker.group(1)
                rest = (marker.group(2) or '').strip()
                split_decimal = re.match(r'^(\d{1,2})\s+(.+)', rest)
                if raw_id.endswith(('.', ',')) and split_decimal:
                    raw_id = f'{raw_id[:-1]}.{split_decimal.group(1)}'
                    rest = split_decimal.group(2).strip()
                if _is_non_english_annex_marker(rest):
                    flush_unnumbered()
                    flush()
                    skip_non_english = True
                    current_section = ''
                    opinion_zone = False
                    continue
                if _is_front_matter_date_marker(raw_id, rest, bool(paragraphs), current_section):
                    continue
                action = _pdf_marker_action(raw_id, rest, current_id)
                if action == 'append_line' and current_id:
                    buf.append(line)
                    continue
                if action == 'append_rest' and current_id:
                    if rest:
                        buf.append(rest)
                    continue
                flush_unnumbered()
                flush()
                current_id = normalize_para_id(raw_id)
                buf = [rest] if rest else []
                continue

            # A numbered marker can sit alone on one line (`4.`) with the
            # paragraph body starting on the next line. In that case even a
            # heading-like sentence ("In its observations ...") is body text.
            if current_id and not buf:
                buf.append(line)
                continue

            previous = ' '.join(buf)
            if _is_pdf_heading(line, previous):
                flush_unnumbered()
                flush()
                if PDF_SECTION_LETTER.match(line):
                    pending_heading = ''
                    continue
                current_section = line.rstrip('.')
                pending_heading = current_section
                opinion_zone = opinion_zone or _is_opinion_section(current_section)
                if _is_opinion_section(current_section) or _is_namespaced_continuation_section(current_section, opinion_zone):
                    unnumbered_counter = 0
                continue

            if current_id:
                buf.append(line)
            elif _is_opinion_section(current_section) or _is_namespaced_continuation_section(current_section, opinion_zone):
                if not unnumbered_buf and _section_needs_continuation(current_section):
                    current_section = f'{current_section} {line}'.strip()
                    pending_heading = current_section
                    continue
                if starts_new_unnumbered_paragraph(line):
                    flush_unnumbered()
                unnumbered_buf.append(line)
            else:
                # Descriptive heading following "A." / "B." style markers.
                if pending_heading == '' and _is_pdf_heading(line):
                    flush_unnumbered()
                    current_section = line.rstrip('.')
                    opinion_zone = opinion_zone or _is_opinion_section(current_section)

    flush_unnumbered()
    flush()
    paragraphs = repair_paragraph_id_sequence(paragraphs)
    return apply_paragraph_namespaces(paragraphs)


def _load_ocr_pages(doc_id: str | None) -> list[str]:
    meta = _load_ocr_meta(doc_id)
    if not meta:
        return []
    meta_path = Path(meta['_metaPath'])
    doc_text = meta_path.parent / (meta.get('textPath') or 'document.txt')
    pages = []
    for page in meta.get('pages') or []:
        page_path = meta_path.parent / (page.get('textPath') or '')
        if page_path.exists():
            pages.append(page_path.read_text())
    if pages:
        return pages
    if doc_text.exists():
        text = doc_text.read_text()
        return [text] if text.strip() else []
    return []


def _load_ocr_meta(doc_id: str | None) -> dict | None:
    if not doc_id:
        return None
    candidates = sorted(OCR_DIR.glob(f'*/{doc_id}/ocr.json'))
    for meta_path in candidates:
        try:
            meta = json.loads(meta_path.read_text())
        except Exception:
            continue
        if meta.get('ocrStatus') not in ('pass', 'review'):
            continue
        meta['_metaPath'] = str(meta_path)
        return meta
    return None


def extract_pdf_paragraphs(path: Path, doc_id: str | None = None) -> list[dict]:
    """Extract numbered jurisprudence paragraphs from a PDF.

    `clean_extract.extract_paragraphs` is tuned for GC/SP reports and only
    recognises top-level `1.` paragraph IDs. Jurisprudence decisions heavily
    use decimal IDs (`1.1`, `2.4`, ...), so we reuse its page cleaning but run
    a jurisprudence-specific line parser here.
    """
    sys.path.insert(0, str(ROOT))
    from clean_extract import _clean_page_text

    doc = fitz.open(path)
    try:
        pages = []
        for page in doc:
            cleaned = _clean_page_text(page)
            pages.append(cleaned if cleaned.strip() else page.get_text())
    finally:
        doc.close()

    paragraphs = _parse_pdf_text_pages(pages)
    if paragraphs:
        return paragraphs
    unnumbered = extract_pdf_unnumbered_decision(pages)
    if unnumbered:
        return unnumbered
    ocr_pages = _load_ocr_pages(doc_id)
    if not ocr_pages:
        return []
    paragraphs = _parse_pdf_text_pages(ocr_pages)
    if paragraphs:
        return paragraphs
    return extract_pdf_unnumbered_decision(ocr_pages)


def extract_pdf_unnumbered_decision(pages: list[str]) -> list[dict]:
    """Fallback for short PDF decisions that have no numbered paragraphs."""
    text = '\n'.join(pages)
    lines = [re.sub(r'\s+', ' ', x.strip()) for x in text.splitlines()]
    lines = [x for x in lines if x]

    body_lines = []
    collecting = False
    for line in lines:
        low = line.lower()
        if low.startswith(('at its meeting', 'the committee', 'decides to discontinue')):
            collecting = True
        if not collecting:
            continue
        if _is_signature_line(line):
            break
        if line.startswith('*') or low.startswith(('united nations', 'international covenant', 'distr.:', 'original:')):
            continue
        body_lines.append(line)

    body = _clean_extracted_text(' '.join(body_lines))
    if len(body) < 80:
        chunks = [
            _clean_extracted_text(c)
            for c in re.split(r'\n\s*\n+', text)
            if len(_clean_extracted_text(c)) >= 20
        ]
        skip_prefixes = (
            'submitted by:', 'alleged victim:', 'state party:',
            'date of decision', 'articles of covenant:',
        )
        body_chunks = []
        collecting_old = False
        for chunk in chunks:
            low = chunk.lower()
            intro_match = re.search(
                r'\b(?:the author of the communication|before considering a communication|the human rights committee)\b',
                chunk,
                re.IGNORECASE,
            )
            if low.startswith(skip_prefixes) and not intro_match:
                continue
            if intro_match:
                chunk = chunk[intro_match.start():]
                low = chunk.lower()
            if (
                low.startswith('the author of the communication')
                or low.startswith('before considering a communication')
                or low.startswith('the human rights committee')
            ):
                collecting_old = True
            if collecting_old:
                body_chunks.append(chunk)
        body = _clean_extracted_text(' '.join(body_chunks))
    if len(body) < 80:
        return []
    if not (
        body.lower().startswith(('at its meeting', 'the committee', 'decides to discontinue'))
        or 'decided to discontinue' in body.lower()
        or 'decides to discontinue' in body.lower()
        or 'declares the communication inadmissible' in body.lower()
        or 'communication is inadmissible' in body.lower()
    ):
        return []
    return [{
        'ID': '1.',
        'Section': 'Decision',
        'Labels': [],
        'Text': body,
    }]


def extract_doc_paragraphs(path: Path) -> list[dict]:
    """Convert a legacy Word `.doc` file to DOCX, then reuse the DOCX parser."""
    textutil = shutil.which('textutil')
    if not textutil:
        raise RuntimeError('legacy .doc parsing requires macOS textutil')
    with tempfile.TemporaryDirectory(prefix='jur_doc_') as tmp:
        out = Path(tmp) / f'{path.stem}.docx'
        subprocess.run(
            [textutil, '-convert', 'docx', '-output', str(out), str(path)],
            check=True,
            capture_output=True,
            text=True,
        )
        return extract_docx_paragraphs(out)


# ---------------------------------------------------------------------------
# Catalog + manifest loading
# ---------------------------------------------------------------------------
def load_catalog() -> list[dict]:
    return [json.loads(l) for l in CATALOG.read_text().splitlines() if l.strip()]


def load_manifest_index() -> dict[str, list[dict]]:
    """Return symbol → [manifest entries]. We only ingest English versions for
    v1, so callers will filter further."""
    by_symbol: dict[str, list[dict]] = {}
    for line in MANIFEST.read_text().splitlines():
        if not line.strip():
            continue
        m = json.loads(line)
        sym = m.get('symbol', '').strip()
        if sym:
            by_symbol.setdefault(sym, []).append(m)
    return by_symbol


def manifest_file_path(entry: dict) -> Path | None:
    rel = (entry.get('file_path') or '').strip()
    if not rel:
        return None
    p = Path(rel)
    if p.is_absolute():
        return p
    parts = p.parts
    if len(parts) >= 2 and parts[0] == 'output' and parts[1] == 'ohchr_jurisprudence':
        return JURIS_SRC / Path(*parts[2:])
    return JURIS_SRC / p


def healthy_manifest_entries(entries: list[dict]) -> list[dict]:
    """English files that exist locally and are non-empty."""
    healthy = []
    for entry in entries:
        if entry.get('language') != 'en':
            continue
        path = manifest_file_path(entry)
        if not path or not path.exists():
            continue
        if path.stat().st_size <= 0:
            continue
        healthy.append(entry)
    return healthy


def choose_english_entry(entries: list[dict]) -> dict | None:
    """Prefer parse-friendly English files, with deterministic fallback."""
    candidates = healthy_manifest_entries(entries)
    if not candidates:
        return None
    fmt_rank = {'docx': 0, 'pdf': 1, 'doc': 2, 'html': 3, 'other': 4}
    candidates.sort(key=lambda m: (
        fmt_rank.get((m.get('format') or '').lower(), 9),
        -(int(m.get('content_length') or 0)),
        m.get('downloaded_at') or '',
    ))
    return candidates[0]


def parse_year_from_symbol(sym: str) -> int | None:
    m = re.search(r'/(\d{4})$', sym)
    return int(m.group(1)) if m else None


def parse_year_from_date(raw: str) -> int | None:
    m = re.search(r'\b(19|20)\d{2}\b', raw or '')
    return int(m.group(0)) if m else None


def shard_id_for(treaty: str, year: int | None) -> str:
    """Sharding rule from JURISPRUDENCE_PLAN.md §3 Tier-2.

    For CCPR (the big one): 2-year buckets. For everyone else: per-treaty
    when small enough, else per-year.  Small treaties (CRPD, CERD, CED)
    fold into a single shard."""
    if treaty in ('CRPD', 'CERD', 'CED'):
        return 'jur_small_treaties'
    if not year:
        return f'jur_{treaty}_unknown'
    if treaty == 'CCPR':
        bucket_start = (year // 2) * 2
        return f'jur_CCPR_{bucket_start}-{bucket_start + 1}'
    if treaty == 'CAT':
        bucket_start = (year // 5) * 5
        return f'jur_CAT_{bucket_start}-{bucket_start + 4}'
    # CRC, CESCR, CEDAW — one shard per treaty
    return f'jur_{treaty}'


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------
def ingest_one(catalog_record: dict, manifest_entries: list[dict]) -> dict | None:
    """Process a single case. Returns the Tier-1 metadata record or None on failure."""
    sym = catalog_record.get('symbol_no', '').strip()
    if not sym:
        return None
    if sym in EXCLUDED_SYMBOLS:
        print(f'    [skip] {sym}: {EXCLUDED_SYMBOLS[sym]}')
        return None
    doc_id = slug(sym)

    # Pick the English manifest entry — prefer DOCX, fall back to PDF, then DOC.
    chosen = choose_english_entry(manifest_entries)
    if not chosen:
        return None
    fmt = (chosen.get('format') or '').lower()
    abs_path = manifest_file_path(chosen)
    if not abs_path or not abs_path.exists():
        return None

    # Extract paragraphs
    try:
        if fmt == 'docx':
            paragraphs = extract_docx_paragraphs(abs_path)
        elif fmt == 'pdf':
            paragraphs = extract_pdf_paragraphs(abs_path, doc_id=doc_id)
        elif fmt == 'doc':
            paragraphs = extract_doc_paragraphs(abs_path)
        else:
            # Unknown extension — skip in v1, but record remains in catalog.
            return None
    except Exception as e:
        print(f'    [warn] {sym}: extraction error — {e}')
        return None

    if not paragraphs:
        return None

    # Apply concerned-group labels
    case_label_set: set[str] = set()
    n_lbl_paragraphs = 0
    for p in paragraphs:
        labels = label_paragraph(p['Text'])
        p['Labels'] = labels
        if labels:
            n_lbl_paragraphs += 1
            case_label_set.update(labels)

    # Outcome classification — title fallbacks to body text (first 12k chars)
    body_blob = ' '.join(p['Text'] for p in paragraphs)
    outcome = classify_outcome(catalog_record.get('title', ''), body_blob)

    ocr_meta = _load_ocr_meta(doc_id) if fmt == 'pdf' else None
    source_format = fmt
    if fmt == 'pdf' and ocr_meta:
        source_format = 'pdf_ocr'
        for p in paragraphs:
            p['SourceFormat'] = 'pdf_ocr'
            p['OcrStatus'] = ocr_meta.get('ocrStatus')
            p['OcrMeanConf'] = ocr_meta.get('meanConf')
            p['OcrLowConfRatio'] = ocr_meta.get('lowConfRatio')

    # docId + paths
    out_para_path = OUT_DIR_PARAGRAPHS / f'{doc_id}.json'
    out_para_path.parent.mkdir(parents=True, exist_ok=True)
    out_para_path.write_text(json.dumps(paragraphs, ensure_ascii=False, indent=2))

    # Year
    year = parse_year_from_symbol(sym)
    raw_date = catalog_record.get('submitted_date', '').strip()
    adoption_year = parse_year_from_date(raw_date)

    # Tier-1 record
    today = date.today().isoformat()
    info = {
        'docId': doc_id,
        'type': 'jur',
        'name': catalog_record.get('title', '').strip() or sym,
        'nameShort': catalog_record.get('title', '').strip() or sym,
        'signature': sym,
        'committee': catalog_record.get('treaty', ''),
        'committees': [catalog_record.get('treaty', '')] if catalog_record.get('treaty') else [],
        'treaty': catalog_record.get('treaty', ''),
        'symbol': sym,
        'country': catalog_record.get('country', '').strip(),
        'year': year,
        'communicationYear': year,
        'adoptionYear': adoption_year,
        'title': catalog_record.get('title', '').strip(),
        'outcome': outcome,
        'submittedDate': raw_date,
        'adoptionDate': raw_date,
        'languages': ['en'],
        'link': catalog_record.get('download_page_url', '').strip(),
        'sourceFile': f'json_jurisprudence/{doc_id}.json',
        'sourceFormat': source_format,
        'shardId': shard_id_for(catalog_record.get('treaty', ''), year),
        'paragraphCount': len(paragraphs),
        'wordCount': sum(len(p['Text'].split()) for p in paragraphs),
        'labelCount': sum(len(p['Labels']) for p in paragraphs),
        'caseLabels': sorted(case_label_set),
        'firstAddedAt': today,
        'lastVerifiedAt': today,
    }
    if ocr_meta:
        info.update({
            'ocrStatus': ocr_meta.get('ocrStatus'),
            'ocrMeanConf': ocr_meta.get('meanConf'),
            'ocrLowConfRatio': ocr_meta.get('lowConfRatio'),
            'ocrPageCount': ocr_meta.get('pageCount'),
            'ocrWordCount': ocr_meta.get('wordCount'),
            'ocrWarnings': ocr_meta.get('warnings') or [],
        })
    return info


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument('--treaty', help='Restrict to one treaty body (e.g. CRPD)')
    ap.add_argument('--limit', type=int, help='Cap number of cases (for testing)')
    ap.add_argument('--all', action='store_true', help='Ingest all treaties')
    args = ap.parse_args()

    if not args.treaty and not args.all:
        ap.print_help()
        return 1

    catalog = load_catalog()
    manifest = load_manifest_index()

    if args.treaty:
        catalog = [r for r in catalog if r.get('treaty', '').upper() == args.treaty.upper()]
    before_mismatch = len(catalog)
    catalog = [
        r for r in catalog
        if symbol_matches_treaty(r.get('symbol_no', ''), r.get('treaty', ''))
    ]
    skipped_mismatch = before_mismatch - len(catalog)
    if args.limit:
        catalog = catalog[: args.limit]

    print(f'Ingesting {len(catalog)} jurisprudence cases'
          f'{" (treaty=" + args.treaty + ")" if args.treaty else ""}…')

    # Load any existing Tier-1 catalog so we can merge.
    existing_records = []
    if OUT_INFO.exists():
        try:
            existing_records = json.loads(OUT_INFO.read_text())
        except Exception:
            existing_records = []
    if args.treaty:
        existing_records = [
            r for r in existing_records
            if (r.get('treaty') or '').upper() != args.treaty.upper()
        ]
    by_doc_id = {r['docId']: r for r in existing_records}

    n_ok = 0
    n_skip = 0
    n_fail = 0

    for i, rec in enumerate(catalog, 1):
        sym = rec.get('symbol_no', '').strip()
        if not sym:
            n_skip += 1
            continue
        manifest_entries = manifest.get(sym, [])
        if not manifest_entries:
            n_skip += 1
            continue
        info = ingest_one(rec, manifest_entries)
        if info is None:
            n_fail += 1
            if i % 25 == 0 or i == len(catalog):
                print(f'  [{i:4d}/{len(catalog)}] ok={n_ok} skip={n_skip} fail={n_fail}')
            continue
        by_doc_id[info['docId']] = info
        n_ok += 1
        if i % 25 == 0 or i == len(catalog):
            print(f'  [{i:4d}/{len(catalog)}] ok={n_ok} skip={n_skip} fail={n_fail}')

    # Write Tier-1 catalog
    OUT_INFO.parent.mkdir(parents=True, exist_ok=True)
    OUT_INFO.write_text(json.dumps(list(by_doc_id.values()), ensure_ascii=False, indent=2))

    print(f'\n=== Ingestion summary ===')
    print(f'  Successful:      {n_ok}')
    print(f'  Skipped (no en): {n_skip}')
    print(f'  Skipped mismatch:{skipped_mismatch}')
    print(f'  Failed:          {n_fail}')
    print(f'  Output paragraphs: {OUT_DIR_PARAGRAPHS}/<docId>.json  ({n_ok} files)')
    print(f'  Output catalog:    {OUT_INFO}  ({len(by_doc_id)} total records)')

    if n_ok:
        # Per-outcome breakdown
        from collections import Counter
        outcomes = Counter(r['outcome'] for r in by_doc_id.values()
                          if not args.treaty or r['treaty'].upper() == args.treaty.upper())
        print(f'\n  Outcome distribution:')
        for o, n in outcomes.most_common():
            print(f'    {o:24s} {n}')

    return 0


if __name__ == '__main__':
    sys.exit(main())
